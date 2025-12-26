import streamlit as st
import datetime
import pandas as pd
import requests
import json
import os
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

# --- KONFIGURASI HALAMAN ---
st.set_page_config(page_title="Sistem Administrasi Guru Sultan AI", layout="wide", page_icon="🏫")

# ==========================================
# 🛑 TEMPEL API KEY ANDA DI BAWAH INI (DI DALAM TANDA KUTIP)
# ==========================================
API_KEY_MANUAL = "AIzaSy..." # <--- Hapus tulisan ini dan tempel API Key Anda di sini

# ==========================================
# 1. SISTEM PELACAKAN (TRACKER)
# ==========================================
STATS_FILE = "daily_stats.csv"

def get_jakarta_time():
    utc_now = datetime.datetime.utcnow()
    jakarta_time = utc_now + datetime.timedelta(hours=7)
    return jakarta_time

def manage_stats(action=None):
    now_jakarta = get_jakarta_time()
    today_str = now_jakarta.strftime("%Y-%m-%d")
    
    if not os.path.exists(STATS_FILE):
        df = pd.DataFrame(columns=["date", "login_count", "gen_count"])
        df.to_csv(STATS_FILE, index=False)
    
    df = pd.read_csv(STATS_FILE)
    
    if today_str not in df['date'].values:
        new_row = pd.DataFrame({"date": [today_str], "login_count": [0], "gen_count": [0]})
        df = pd.concat([df, new_row], ignore_index=True)
    
    if action == 'login':
        df.loc[df['date'] == today_str, 'login_count'] += 1
    elif action == 'generate':
        df.loc[df['date'] == today_str, 'gen_count'] += 1
        
    df.to_csv(STATS_FILE, index=False)
    today_data = df.loc[df['date'] == today_str].iloc[0]
    return today_data['login_count'], today_data['gen_count'], df

# ==========================================
# 2. CSS SKEUOMORPHISM & STYLE
# ==========================================
st.markdown("""
<style>
    .stApp { background-color: #e0e5ec; color: #4d4d4d; font-family: 'Segoe UI', sans-serif; }
    
    /* Card Style */
    .skeuo-card {
        border-radius: 20px;
        background: #e0e5ec;
        box-shadow:  9px 9px 16px rgb(163,177,198,0.6), -9px -9px 16px rgba(255,255,255, 0.5);
        padding: 25px; margin-bottom: 20px;
        border: 1px solid rgba(255,255,255,0.2);
    }
    
    /* Input Style */
    .stTextInput>div>div>input, .stTextArea>div>div>textarea, .stSelectbox>div>div>div, .stNumberInput>div>div>input, .stDateInput>div>div>input {
        background-color: #e0e5ec !important;
        border-radius: 10px; border: none;
        box-shadow: inset 4px 4px 8px #bebebe, inset -4px -4px 8px #ffffff !important;
        color: #333 !important; padding: 10px;
    }
    
    /* Multiselect Tag Style */
    .stMultiSelect span[data-baseweb="tag"] {
        background-color: #e0e5ec !important;
        border: 1px solid #ccc;
        border-radius: 10px;
    }

    /* Button Style */
    .stButton>button {
        width: 100%; border: none; outline: none; border-radius: 12px;
        background: linear-gradient(145deg, #f0f0f3, #cacaca);
        box-shadow:  6px 6px 12px #bebebe, -6px -6px 12px #ffffff;
        color: #0d47a1; font-weight: bold; padding: 10px 20px;
        transition: all 0.2s ease-in-out;
    }
    .stButton>button:active {
        background: #e0e5ec;
        box-shadow: inset 4px 4px 8px #bebebe, inset -4px -4px 8px #ffffff;
        transform: translateY(2px);
    }
    
    /* Header & Footer */
    .header-container {
        display: flex; justify-content: space-between; align-items: center;
        background: #e0e5ec; padding: 15px 20px; border-radius: 15px;
        box-shadow: 5px 5px 10px #bebebe, -5px -5px 10px #ffffff; margin-bottom: 25px;
    }
    .live-clock {
        background: #e0e5ec; padding: 5px 15px; border-radius: 8px;
        box-shadow: inset 3px 3px 6px #bebebe, inset -3px -3px 6px #ffffff;
        color: #0d47a1; font-weight: bold; font-size: 14px; text-align: right;
    }
    
    /* FOOTER FIXED */
    .footer {
        position: fixed; left: 0; bottom: 0; width: 100%;
        background-color: #e0e5ec; color: #555; text-align: center;
        padding: 10px; font-weight: bold; box-shadow: 0px -4px 10px rgba(0,0,0,0.1); z-index: 9999;
        font-size: 14px;
    }
    
    h3 { color: #0d47a1; font-weight: bold; margin-bottom: 15px; }
    h4 { color: #333; font-weight: bold; margin-bottom: 10px; font-size: 16px; }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 3. HEADER & JAM (JS)
# ==========================================
def render_header():
    st.markdown("""
        <div class="header-container">
            <div style="width: 65%; font-family: 'Courier New', monospace; font-weight: bold; color: #2c3e50; font-size: 16px;">
                <marquee direction="left" scrollamount="6">🚀 SISTEM ADMINISTRASI GURU TERPADU - SD MUHAMMADIYAH 8 TULANGAN 🚀</marquee>
            </div>
            <div id="clock" class="live-clock">Loading...</div>
        </div>
        <script>
            function updateTime() {
                const now = new Date();
                const optionsDate = { weekday: 'long', year: 'numeric', month: 'long', day: 'numeric' };
                document.getElementById('clock').innerHTML = now.toLocaleDateString('id-ID', optionsDate) + '<br>' + now.toLocaleTimeString('id-ID', {hour:'2-digit', minute:'2-digit', second:'2-digit'}) + ' WIB';
            }
            setInterval(updateTime, 1000); updateTime();
        </script>
        
        <div class="footer">Aplikasi by Haris Adz Dzimari &copy; 2025</div>
    """, unsafe_allow_html=True)

# ==========================================
# 4. FUNGSI LOGIKA (AI DIRECT API - ANTI ERROR)
# ==========================================
def tanya_gemini(api_key, prompt):
    if not api_key or "AIza" not in api_key: 
        return "⚠️ API Key belum diisi atau salah format! Cek baris 16 di kode app.py"
    
    headers = {'Content-Type': 'application/json'}
    data = {"contents": [{"parts": [{"text": prompt}]}]}
    
    # DAFTAR MODEL (Akan dicoba satu per satu sampai berhasil)
    models_to_try = [
        "gemini-1.5-flash",        # Opsi 1: Paling Cepat
        "gemini-1.5-flash-latest", # Opsi 2: Alternatif Flash
        "gemini-pro"               # Opsi 3: Paling Stabil (Legacy)
    ]
    
    last_error = ""
    
    for model in models_to_try:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
        try:
            response = requests.post(url, headers=headers, json=data)
            if response.status_code == 200:
                # BERHASIL! Langsung kembalikan teks
                return response.json()['candidates'][0]['content']['parts'][0]['text']
            else:
                last_error = f"Model {model} Gagal ({response.status_code}). Mencoba model lain..."
                continue # Coba model berikutnya
        except Exception as e:
            last_error = str(e)
            continue

    # Jika semua model gagal
    return f"❌ Gagal Generate. Semua model sibuk/error. Pesan terakhir: {last_error}"

# FUNGSI EXPORT DOCX SEDERHANA
def create_simple_docx(title, content, sekolah):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Sekolah: {sekolah}")
    doc.add_paragraph("_"*50)
    doc.add_paragraph(content)
    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# FUNGSI EXPORT DOCX MODUL AJAR
def create_modul_docx(data):
    doc = Document()
    for s in doc.sections: s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.54)

    if data['logo']:
        try:
            doc.add_picture(data['logo'], width=Inches(1.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except: pass

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"MODUL AJAR KURIKULUM MERDEKA\n{data['sekolah']}\n")
    r.bold = True; r.font.size = Pt(14)
    p.add_run(f"{data['alamat']}").font.size = Pt(10)
    doc.add_paragraph("_"*85)

    # I. INFO UMUM
    doc.add_heading('I. INFORMASI UMUM', 1)
    table = doc.add_table(rows=0, cols=2); table.style = 'Table Grid'
    info = [
        ("Penyusun", data['guru']), 
        ("Tahun", str(data['tanggal'].year)), 
        ("Jenjang / Kelas", f"{data['kelas']} ({data['fase']})"), 
        ("Mata Pelajaran", data['mapel']), 
        ("Topik / Bab", data['topik']), 
        ("Alokasi Waktu", data['alokasi']),
        ("Model Pembelajaran", data['model'])
    ]
    for k,v in info:
        row = table.add_row()
        row.cells[0].text = k; row.cells[0].paragraphs[0].runs[0].bold = True; row.cells[1].text = v
    
    doc.add_paragraph(f"\nCapaian Pembelajaran (CP): {data['cp']}")
    doc.add_paragraph(f"Profil Pelajar: {', '.join(data['profil'])}")

    # II. INTI
    doc.add_heading('II. KOMPONEN INTI', 1)
    doc.add_heading('A. Tujuan Pembelajaran', 2); doc.add_paragraph(data['tujuan'])
    doc.add_heading('B. Pertanyaan Pemantik', 2); doc.add_paragraph(data['pemantik'])
    
    doc.add_heading('C. Diferensiasi', 2)
    doc.add_paragraph(f"Remedial: {data['remedial']}")
    doc.add_paragraph(f"Pengayaan: {data['pengayaan']}")

    # III. KEGIATAN
    doc.add_heading('III. KEGIATAN PEMBELAJARAN', 1)
    doc.add_heading('1. Ringkasan Materi', 2); doc.add_paragraph(data['bahan'])
    doc.add_heading('2. Langkah LKPD', 2); doc.add_paragraph(data['lkpd'])
    doc.add_heading('3. Media Ajar', 2); doc.add_paragraph(data['media'])

    # IV. EVALUASI
    doc.add_heading('IV. EVALUASI', 1)
    doc.add_heading('A. Soal Latihan', 2); doc.add_paragraph(data['soal'])
    doc.add_heading('B. Kunci Jawaban', 2); doc.add_paragraph(data['kunci'])

    # V. ASESMEN & RUBRIK
    doc.add_heading('V. ASESMEN', 1)
    doc.add_paragraph(f"Teknik Penilaian: {', '.join(data['teknik_nilai'])}")
    
    doc.add_heading('Rubrik Penilaian Sikap (Profil Pelajar)', 2)
    t_rubrik = doc.add_table(rows=1, cols=5); t_rubrik.style = 'Table Grid'
    rh = t_rubrik.rows[0].cells
    rh[0].text="Dimensi"; rh[1].text="Membudaya (4)"; rh[2].text="Berkembang (3)"; rh[3].text="Mulai Terlihat (2)"; rh[4].text="Belum Terlihat (1)"
    for dim in data['profil']:
        rr = t_rubrik.add_row().cells
        rr[0].text = dim
        rr[1].text = "Sangat konsisten menunjukkan sikap ini."
        rr[2].text = "Konsisten menunjukkan sikap ini."
        rr[3].text = "Mulai menunjukkan sikap ini."
        rr[4].text = "Belum menunjukkan sikap ini."

    # VI. DAFTAR HADIR
    doc.add_page_break(); doc.add_heading('VI. DAFTAR HADIR SISWA', 1)
    doc.add_paragraph(f"Kelas: {data['kelas']} | Tanggal: {data['tanggal'].strftime('%d-%m-%Y')}")
    
    t_absen = doc.add_table(rows=1, cols=5); t_absen.style = 'Table Grid'
    hdr = t_absen.rows[0].cells; hdr[0].text="No"; hdr[1].text="Nama Siswa"; hdr[2].text="Hadir"; hdr[3].text="Sakit/Izin"; hdr[4].text="Ket"
    siswa = data['siswa_list'] if data['siswa_list'] else [""]*25
    for i, nm in enumerate(siswa):
        r = t_absen.add_row().cells; r[0].text=str(i+1); r[1].text=nm.strip()

    # VII. LAMPIRAN
    doc.add_paragraph("\n"); doc.add_heading('VII. LAMPIRAN', 1)
    doc.add_heading('Daftar Pustaka', 2); doc.add_paragraph(data['pustaka'])
    doc.add_heading('Glosarium', 2); doc.add_paragraph(data['glosarium'])
    
    # VIII. REFLEKSI
    doc.add_heading('Refleksi', 1)
    doc.add_heading('Refleksi Guru', 2); doc.add_paragraph(data['ref_guru'])
    doc.add_heading('Refleksi Siswa', 2); doc.add_paragraph(data['ref_siswa'])

    # TTD
    doc.add_paragraph("\n\n")
    ttd = doc.add_table(rows=1, cols=2)
    ttd.rows[0].cells[0].text = f"Mengetahui,\nKepala Sekolah\n\n\n( {data['kepsek']} )"
    ttd.rows[0].cells[1].text = f"Sidoarjo, {data['tanggal'].strftime('%d %B %Y')}\nGuru Kelas\n\n\n( {data['guru']} )"

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf

def create_pdf(data):
    pdf = FPDF(); pdf.add_page(); pdf.set_font("Arial", size=11)
    def safe(txt): return txt.encode('latin-1', 'replace').decode('latin-1')
    
    pdf.set_font("Arial", 'B', 14); pdf.cell(0,10, "MODUL AJAR", 0, 1, 'C'); pdf.set_font("Arial", size=11)
    pdf.cell(0, 8, safe(f"Sekolah: {data['sekolah']}"), ln=True)
    pdf.cell(0, 8, safe(f"Guru: {data['guru']}"), ln=True)
    pdf.ln(5)
    pdf.set_font("Arial", 'B', 11); pdf.cell(0, 8, "TUJUAN & TEKNIK PENILAIAN:", ln=True)
    pdf.set_font("Arial", size=11); pdf.multi_cell(0, 6, safe(data['tujuan'])); pdf.ln(3)
    pdf.multi_cell(0, 6, safe(f"Teknik Penilaian: {', '.join(data['teknik_nilai'])}")); pdf.ln(3)
    return pdf.output(dest='S').encode('latin-1', 'replace')

# ==========================================
# 5. HALAMAN FITUR (MODUL, ATP, PROTA)
# ==========================================

def menu_modul_ajar(api_key, nama_sekolah, alamat_sekolah, kepsek, uploaded_logo):
    st.subheader("📂 Generator Modul Ajar")
    
    # 1. Identitas
    t1, t2, t3, t4, t5, t6, t7 = st.tabs(["1️⃣ Identitas", "2️⃣ Inti", "3️⃣ Bahan & LKPD", "4️⃣ Evaluasi", "5️⃣ Asesmen", "6️⃣ Glosarium", "7️⃣ Refleksi"])

    with t1:
        st.markdown("<div class='skeuo-card'>", unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1: 
            nama_guru = st.text_input("Nama Guru", value=st.session_state.get('nama_guru', ''), placeholder="Nama Lengkap")
            tanggal = st.date_input("Tanggal")
            mapel = st.text_input("Mapel", value=st.session_state.get('mapel', ''), placeholder="Informatika")
        with c2: 
            fase = st.selectbox("Fase", ["Fase A (Kls 1-2)", "Fase B (Kls 3-4)", "Fase C (Kls 5-6)"])
            kelas = st.selectbox("Kelas", ["1","2","3","4","5","6"])
            alokasi = st.text_input("Alokasi Waktu", value="2 JP (2 x 35 Menit)")
        cp = st.text_area("Capaian Pembelajaran (CP):", height=100)
        st.markdown("</div>", unsafe_allow_html=True)

    with t2:
        st.markdown("<div class='skeuo-card'>", unsafe_allow_html=True)
        col_inti_1, col_inti_2 = st.columns(2)
        with col_inti_1:
            st.markdown("### 📚 Materi & Tujuan")
            topik = st.text_input("Topik / Bab", value=st.session_state.get('topik', ''))
            model = st.selectbox("Model Pembelajaran", ["Problem-Based Learning (PBL)", "Project-Based Learning (PjBL)", "Discovery Learning (DL)", "Inquiry Learning (IL)"])
            if st.button("✨ Bantu Buat Tujuan"):
                with st.spinner("AI Bekerja..."):
                    manage_stats('generate') 
                    p = f"Buatkan tujuan pembelajaran (TP) dan pertanyaan pemantik untuk mapel {mapel} topik {topik} fase {fase} model {model}."
                    st.session_state['tujuan_ai'] = tanya_gemini(api_key, p)
                    st.rerun()
            tujuan = st.text_area("Tujuan Pembelajaran (TP)", value=st.session_state.get('tujuan_ai', ''), height=150)
            pemantik = st.text_input("Pertanyaan Pemantik", placeholder="Mengapa kita perlu...?")

        with col_inti_2:
            st.markdown("### 👤 Profil & Diferensiasi")
            st.write("8 Dimensi Profil:")
            profil_opsi = ["Keimanan & Ketakwaan", "Kewargaan", "Penalaran Kritis", "Kreativitas", "Kolaborasi", "Kemandirian", "Kesehatan (Fisik & Mental)", "Komunikasi"]
            profil = st.multiselect("Pilih Dimensi", profil_opsi, default=["Penalaran Kritis", "Kreativitas"], label_visibility="collapsed")
            st.divider()
            remedial = st.text_area("Remedial:", value="Pendampingan individu dan penyederhanaan materi.", height=80)
            pengayaan = st.text_area("Pengayaan:", value="Tugas proyek tambahan atau tutor sebaya.", height=80)
        st.markdown("</div>", unsafe_allow_html=True)

    with t3:
        st.markdown("<div class='skeuo-card'>", unsafe_allow_html=True)
        if st.button("✨ Auto Materi & LKPD"):
            with st.spinner("Menyusun..."):
                manage_stats('generate') 
                st.session_state['materi_ai'] = tanya_gemini(api_key, f"Ringkasan materi {topik} SD kelas {kelas}.")
                st.session_state['lkpd_ai'] = tanya_gemini(api_key, f"Buatkan petunjuk LKPD aktivitas siswa topik {topik}.")
                st.session_state['media_ai'] = tanya_gemini(api_key, f"List media ajar untuk topik {topik}.")
                
                hist_data = {'waktu': datetime.datetime.now().strftime("%H:%M"), 'topik': topik, 'data': {'tujuan_ai': st.session_state.get('tujuan_ai', ''), 'materi_ai': st.session_state['materi_ai'], 'lkpd_ai': st.session_state['lkpd_ai'], 'media_ai': st.session_state['media_ai'], 'topik': topik, 'mapel': mapel}}
                if 'history' not in st.session_state: st.session_state['history'] = []
                st.session_state['history'].append(hist_data)
                st.rerun()

        st.markdown("#### 📖 Ringkasan Bahan Ajar")
        bahan = st.text_area("Materi:", value=st.session_state.get('materi_ai', ''), height=200)
        st.divider()
        st.markdown("#### 📝 Desain LKPD & Media")
        lkpd = st.text_area("Petunjuk LKPD:", value=st.session_state.get('lkpd_ai', ''), height=200)
        media = st.text_area("Media Ajar:", value=st.session_state.get('media_ai', ''), height=80)
        st.markdown("</div>", unsafe_allow_html=True)

    with t4:
        st.markdown("<div class='skeuo-card'>", unsafe_allow_html=True)
        if st.button("✨ Auto Soal & Kunci"):
             with st.spinner("Membuat Soal..."):
                 manage_stats('generate') 
                 st.session_state['soal_ai'] = tanya_gemini(api_key, f"Buatkan 5 Soal Essay HOTS tentang {topik}.")
                 st.session_state['kunci_ai'] = tanya_gemini(api_key, f"Buatkan Kunci Jawaban untuk soal essay topik {topik}.")
                 st.rerun()
        c_ev1, c_ev2 = st.columns(2)
        with c_ev1:
            st.markdown("#### ❓ Soal Latihan")
            soal = st.text_area("Daftar Soal:", value=st.session_state.get('soal_ai', ''), height=250)
        with c_ev2:
            st.markdown("#### 🔑 Kunci Jawaban")
            kunci = st.text_area("Kunci Jawaban:", value=st.session_state.get('kunci_ai', ''), height=250)
        st.markdown("</div>", unsafe_allow_html=True)

    with t5:
        st.markdown("<div class='skeuo-card'>", unsafe_allow_html=True)
        st.write("Teknik Penilaian:")
        teknik_nilai = st.multiselect("Pilih Teknik", ["Tes Tulis", "Lisan", "Observasi", "Unjuk Kerja", "Portofolio", "Proyek"], default=["Tes Tulis", "Observasi"], label_visibility="collapsed")
        
        st.divider()
        st.write("Preview Rubrik Sikap (Otomatis berdasarkan Profil di Tab 2):")
        if profil:
            data_rubrik = []
            for p in profil:
                data_rubrik.append({"Dimensi": p, "Skor 4": "Membudaya", "Skor 3": "Berkembang", "Skor 2": "Mulai Terlihat", "Skor 1": "Belum Terlihat"})
            st.table(pd.DataFrame(data_rubrik))
        else: st.warning("⚠️ Belum ada Dimensi Profil yang dipilih di Tab 2.")

        st.divider()
        st.write("Ketik nama siswa (1 nama per baris) untuk mengisi Tabel Absensi otomatis:")
        raw_siswa = st.text_area("Daftar Nama Siswa:", height=150, placeholder="Adi\nBudi\nCici...")
        siswa_list = [x for x in raw_siswa.split('\n') if x.strip()]
        st.write(f"Terdeteksi: **{len(siswa_list)} Siswa**")
        st.markdown("</div>", unsafe_allow_html=True)

    with t6:
        st.markdown("<div class='skeuo-card'>", unsafe_allow_html=True)
        pustaka = st.text_area("📚 Daftar Pustaka:", value=f"1. Buku Paket Kemdikbud Kelas {kelas}", height=100)
        glosarium = st.text_area("🔤 Glosarium:", placeholder="Istilah sulit...", height=100)
        st.markdown("</div>", unsafe_allow_html=True)

    with t7:
        st.markdown("<div class='skeuo-card'>", unsafe_allow_html=True)
        c_ref1, c_ref2 = st.columns(2)
        with c_ref1:
            st.markdown("#### 👨‍🏫 Refleksi Guru")
            ref_guru = st.text_area("Catatan Guru:", placeholder="Kendala, keberhasilan...", height=150)
        with c_ref2:
            st.markdown("#### 🧒 Refleksi Siswa")
            ref_siswa = st.text_area("Catatan Siswa:", placeholder="Respon siswa...", height=150)
        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='skeuo-card' style='text-align:center;'>", unsafe_allow_html=True)
    st.success("✅ Dokumen Siap Unduh")
    data_export = {
        'logo': uploaded_logo, 'sekolah': nama_sekolah, 'alamat': alamat_sekolah, 'kepsek': kepsek,
        'guru': nama_guru, 'mapel': mapel, 'kelas': kelas, 'fase': fase, 'tanggal': tanggal, 'alokasi': alokasi,
        'cp': cp, 'topik': topik, 'model': model, 'tujuan': tujuan, 'pemantik': pemantik,
        'profil': profil, 'remedial': remedial, 'pengayaan': pengayaan,
        'bahan': bahan, 'lkpd': lkpd, 'media': media, 'soal': soal, 'kunci': kunci,
        'teknik_nilai': teknik_nilai,
        'siswa_list': siswa_list, 'pustaka': pustaka, 'glosarium': glosarium, 
        'ref_guru': ref_guru, 'ref_siswa': ref_siswa
    }
    c_dl1, c_dl2 = st.columns(2)
    with c_dl1:
        st.download_button("📄 DOWNLOAD WORD (.DOCX)", create_modul_docx(data_export), f"Modul_{topik}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with c_dl2:
        st.download_button("📕 DOWNLOAD PDF", create_pdf(data_export), f"Modul_{topik}.pdf", "application/pdf")
    st.markdown("</div>", unsafe_allow_html=True)

def menu_atp(api_key, nama_sekolah):
    st.subheader("🗺️ Generator Alur Tujuan Pembelajaran (ATP)")
    st.markdown("<div class='skeuo-card'>", unsafe_allow_html=True)
    mapel = st.text_input("Mata Pelajaran (ATP)", "IPAS")
    fase = st.selectbox("Fase (ATP)", ["A", "B", "C"])
    kelas = st.selectbox("Kelas (ATP)", ["1", "2", "3", "4", "5", "6"])
    cp_text = st.text_area("Capaian Pembelajaran (CP) yang ingin dipecah:", height=150)
    
    if st.button("✨ Generate ATP"):
        with st.spinner("Merancang ATP..."):
            manage_stats('generate')
            prompt = f"Buatkan Tabel Alur Tujuan Pembelajaran (ATP) untuk Mapel {mapel} Fase {fase} Kelas {kelas}. Dari CP berikut: {cp_text}. Kolom tabel: No, Elemen, Capaian Pembelajaran, Tujuan Pembelajaran, Alokasi Waktu, Profil Pelajar Pancasila."
            result = tanya_gemini(api_key, prompt)
            st.session_state['atp_result'] = result
    
    if 'atp_result' in st.session_state:
        st.markdown(st.session_state['atp_result'])
        st.download_button("📥 Simpan ATP (.docx)", create_simple_docx("ALUR TUJUAN PEMBELAJARAN (ATP)", st.session_state['atp_result'], nama_sekolah), "ATP.docx")
    st.markdown("</div>", unsafe_allow_html=True)

def menu_prota(api_key, nama_sekolah):
    st.subheader("📅 Generator Program Tahunan (Prota)")
    st.markdown("<div class='skeuo-card'>", unsafe_allow_html=True)
    mapel = st.text_input("Mata Pelajaran (Prota)", "Matematika")
    kelas = st.selectbox("Kelas (Prota)", ["1", "2", "3", "4", "5", "6"])
    
    if st.button("✨ Generate Prota"):
        with st.spinner("Menyusun Prota..."):
            manage_stats('generate')
            prompt = f"Buatkan Program Tahunan (Prota) untuk Mapel {mapel} SD Kelas {kelas} Kurikulum Merdeka. Distribusikan materi untuk Semester 1 dan Semester 2. Format Tabel: No, Semester, Bab/Topik, Tujuan Pembelajaran Ringkas, Alokasi Waktu."
            result = tanya_gemini(api_key, prompt)
            st.session_state['prota_result'] = result
    
    if 'prota_result' in st.session_state:
        st.markdown(st.session_state['prota_result'])
        st.download_button("📥 Simpan Prota (.docx)", create_simple_docx("PROGRAM TAHUNAN", st.session_state['prota_result'], nama_sekolah), "Prota.docx")
    st.markdown("</div>", unsafe_allow_html=True)

# ==========================================
# 6. APLIKASI UTAMA (MAIN APP)
# ==========================================
def main_app():
    render_header()
    logins, gens, df_stats = manage_stats() 
    utc_now = datetime.datetime.utcnow(); jakarta_time = utc_now + datetime.timedelta(hours=7)
    today_date = jakarta_time.strftime("%d %B %Y"); now_time = jakarta_time.strftime("%H:%M WIB")

    st.markdown("<div class='skeuo-card' style='text-align:center;'><h1 style='color:#0d47a1; margin:0;'>💎 SUPER APP GURU</h1></div>", unsafe_allow_html=True)

    with st.sidebar:
        st.markdown("<div class='skeuo-card' style='text-align:center;'>⚙️ <b>NAVIGASI</b></div>", unsafe_allow_html=True)
        
        # --- MENU PILIHAN ---
        menu = st.radio("Pilih Alat:", ["📂 Modul Ajar", "🗺️ Generator ATP", "📅 Generator Prota"], index=0)
        
        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown(f"""
        <div style='background:#f0f2f6; padding:10px; border-radius:10px; margin-bottom:15px; text-align:center;'>
            <h4 style='margin:0;'>📊 Statistik Hari Ini</h4>
            <p style='font-size:12px; margin-bottom:5px; color:#555;'>{today_date} | {now_time}</p>
            <p style='margin:0;'>Login: <b>{logins}</b> | Aktivitas: <b>{gens}</b></p>
        </div>
        """, unsafe_allow_html=True)
        
        # Grafik
        if not df_stats.empty:
            st.caption("Tren Aktivitas (7 Hari)")
            st.bar_chart(df_stats.tail(7).set_index('date')['gen_count'])

        # API KEY DARI VARIABEL MANUAL / SECRETS
        if API_KEY_MANUAL != "":
            api_key = API_KEY_MANUAL
        elif "GEMINI_API_KEY" in st.secrets: 
            api_key = st.secrets["GEMINI_API_KEY"]
        else: 
            api_key = ""
        
        if menu == "📂 Modul Ajar":
            if st.button("🔄 Reset Modul"):
                keys = ['tujuan_ai', 'materi_ai', 'lkpd_ai', 'media_ai', 'soal_ai', 'kunci_ai']; 
                for k in keys: 
                    if k in st.session_state: del st.session_state[k]
                st.rerun()

        st.divider()
        st.write("<b>Identitas Sekolah:</b>", unsafe_allow_html=True)
        uploaded_logo = st.file_uploader("Upload Logo", type=['png', 'jpg', 'jpeg'])
        nama_sekolah = st.text_input("Sekolah", value="SD MUHAMMADIYAH 8 TULANGAN")
        alamat_sekolah = st.text_area("Alamat", value="Jl. Raya Kenongo RT. 02 RW. 01 Tulangan Sidoarjo")
        kepsek = st.text_input("Kepala Sekolah", value="Muhammad Saifudin Zuhri, M.Pd.")
        
        if st.button("Logout"): 
            st.session_state['logged_in'] = False
            st.query_params.clear()
            st.rerun()
            
        # History (Hanya muncul di Modul Ajar)
        if menu == "📂 Modul Ajar":
            st.divider(); st.write("📜 **Riwayat Sesi Ini:**")
            if 'history' not in st.session_state: st.session_state['history'] = []
            if st.session_state['history']:
                for i, h in enumerate(reversed(st.session_state['history'])):
                    with st.expander(f"{h['topik']} ({h['waktu']})"):
                        if st.button("Muat", key=f"load_{i}"): st.session_state.update(h['data']); st.rerun()
            else: st.caption("Kosong")

    # --- RENDER KONTEN BERDASARKAN MENU ---
    if menu == "📂 Modul Ajar":
        menu_modul_ajar(api_key, nama_sekolah, alamat_sekolah, kepsek, uploaded_logo)
    elif menu == "🗺️ Generator ATP":
        menu_atp(api_key, nama_sekolah)
    elif menu == "📅 Generator Prota":
        menu_prota(api_key, nama_sekolah)

# LOGIN & SESSION
if "auth" in st.query_params and st.query_params["auth"] == "true": st.session_state['logged_in'] = True
if 'logged_in' not in st.session_state: st.session_state['logged_in'] = False

if not st.session_state['logged_in']:
    render_header()
    st.markdown("<br><br><div class='skeuo-card' style='max-width:400px; margin:auto; text-align:center;'><h2>🔐 LOGIN</h2><hr></div>", unsafe_allow_html=True)
    c1,c2,c3 = st.columns([1,1,1])
    with c2:
        u = st.text_input("User"); p = st.text_input("Pass", type="password")
        if st.button("MASUK"): 
            if u=="guru" and p=="123": 
                manage_stats('login') 
                st.session_state['logged_in'] = True
                st.query_params["auth"] = "true"
                st.rerun()
            else: st.error("Gagal")
else: 
    main_app()
